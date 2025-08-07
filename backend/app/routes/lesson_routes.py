# backend/app/routes/lesson_routes.py
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from app.utils.rag_utils import build_faiss_index, retrieve_strategies
from app.utils.llm_utils import generate_rule_mapping, personalize_lesson
import json, os
from uuid import uuid4
from docx import Document

router = APIRouter()
LESSON_DIR = "data/lessons"
OUTPUT_DIR = "data/outputs"
PROFILE_DIR = "data/profiles"
RULE_DIR = "data/rules"

# Helper: Extract text from DOCX
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return full_text

@router.post("/upload")
def upload_lesson(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only DOCX files are supported.")
    file_id = str(uuid4())
    file_path = os.path.join(LESSON_DIR, f"{file_id}.docx")
    with open(file_path, "wb") as f:
        f.write(file.file.read())
    return {"message": "Lesson uploaded.", "lesson_id": file_id}

@router.post("/personalize/{lesson_id}/{profile_name}")
def personalize(lesson_id: str, profile_name: str):
    # Load profile
    profile_path = os.path.join(PROFILE_DIR, f"{profile_name}.json")
    if not os.path.exists(profile_path):
        raise HTTPException(status_code=404, detail="Profile not found.")
    with open(profile_path, "r") as f:
        profile_data = json.load(f)

    # Build index for strategy retrieval
    vectorstore = build_faiss_index()
    query = f"Adaptation strategies for {profile_data['disability']} and language {profile_data['native_reading_level']}"
    strategies = retrieve_strategies(query, vectorstore)

    # Auto-fill 'needs' if missing or empty
    if not profile_data.get("needs"):
        extracted_needs = []
        for text in strategies[:3]:
            if "instructions" in text:
                extracted_needs.append("chunked instructions")
            if "visual" in text:
                extracted_needs.append("visuals")
            if "distraction" in text:
                extracted_needs.append("reduced distractions")
            if "simplified" in text or "simple language" in text:
                extracted_needs.append("simplified language")
        profile_data["needs"] = list(set(extracted_needs))

        # Save updated profile
        with open(profile_path, "w") as f:
            json.dump(profile_data, f, indent=2)

    # Generate rule mapping
    rule_mapping = generate_rule_mapping(profile_data, strategies)

    # Save rule mapping
    rule_filename = f"{profile_name}_rules.json"
    with open(os.path.join(RULE_DIR, rule_filename), "w") as f:
        f.write(rule_mapping)

    # Extract lesson text
    lesson_path = os.path.join(LESSON_DIR, f"{lesson_id}.docx")
    if not os.path.exists(lesson_path):
        raise HTTPException(status_code=404, detail="Lesson not found.")
    lesson_text = extract_text_from_docx(lesson_path)

    # Personalize lesson
    modified_lesson = personalize_lesson(lesson_text, rule_mapping)

    # Save modified lesson
    output_filename = f"{lesson_id}_modified.docx"
    output_file_path = os.path.join(OUTPUT_DIR, output_filename)
    doc = Document()
    for line in modified_lesson.split('\n'):
        doc.add_paragraph(line.strip())
    doc.save(output_file_path)

    return {
        "message": "Lesson personalized.",
        "download_url": f"/api/lessons/download/{output_filename}",
        "filename": output_filename
    }

@router.get("/download/{filename}")
def download_lesson(filename: str):
    safe_filename = os.path.basename(filename)
    file_path = os.path.join(OUTPUT_DIR, safe_filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found.")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_filename
    )